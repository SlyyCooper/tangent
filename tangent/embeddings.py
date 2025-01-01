from typing import List, Optional, Union
from pathlib import Path
import json
import os
import uuid
import markdown
import docx
import PyPDF2
from openai import OpenAI
from qdrant_client import QdrantClient
from qdrant_client.http import models as rest

# Import types from types.py
from .types import (
    Document,
    DocumentChunk,
    EmbeddingConfig,
    VectorDBConfig,
    QdrantConfig,
    PineconeConfig,
    CustomVectorDBConfig
)

# Keep the EMBEDDING_DIMENSIONS constant
EMBEDDING_DIMENSIONS = {
    "text-embedding-3-large": 3072,
    "text-embedding-3-small": 1536,
}

# Keep only the EmbeddingManager class
class EmbeddingManager:
    """Manages document loading, embedding creation, and search."""
    
    def __init__(self, config: EmbeddingConfig, client=None):
        self.config = config
        self.client = client or OpenAI()
        
        # Set up vector database
        match self.config.vector_db:
            case QdrantConfig() as qdrant_config:
                self.vector_db = QdrantClient(
                    host=qdrant_config.url,
                    port=qdrant_config.port,
                    api_key=qdrant_config.api_key
                )
            case PineconeConfig() as pinecone_config:
                import pinecone
                pinecone.init(
                    api_key=pinecone_config.api_key,
                    environment=pinecone_config.environment
                )
                self.vector_db = pinecone.Index(pinecone_config.index_name)
            case CustomVectorDBConfig() as custom_config:
                # User needs to provide their own vector_db instance
                self.vector_db = None
        
        # Get vector size for the model
        if config.model not in EMBEDDING_DIMENSIONS:
            test_response = self.client.embeddings.create(
                model=self.config.model,
                input="test"
            )
            self.vector_size = len(test_response.data[0].embedding)
        else:
            self.vector_size = EMBEDDING_DIMENSIONS[config.model]
        
        self._setup_collection()
    
    def _setup_collection(self):
        """Setup the vector collection based on the configured database type."""
        match self.config.vector_db:
            case QdrantConfig():
                self._setup_qdrant_collection()
            case PineconeConfig():
                # Pinecone indexes are created through their dashboard
                pass
            case CustomVectorDBConfig():
                # Custom DB setup should be handled by the user
                pass
    
    def _setup_qdrant_collection(self):
        """Setup Qdrant collection."""
        try:
            collection_info = self.vector_db.get_collection(self.config.vector_db.collection_name)
            if collection_info.config.params.vectors.size != self.vector_size:
                if not self.config.recreate_collection:
                    raise ValueError(
                        f"Existing collection has incorrect vector size "
                        f"({collection_info.config.params.vectors.size} vs {self.vector_size})"
                    )
                self.vector_db.delete_collection(self.config.vector_db.collection_name)
                
        except Exception as e:
            if "Not found" not in str(e):
                raise
        
        # Create or recreate collection
        try:
            self.vector_db.delete_collection(self.config.vector_db.collection_name)
        except:
            pass
            
        self.vector_db.create_collection(
            collection_name=self.config.vector_db.collection_name,
            vectors_config=rest.VectorParams(
                size=self.vector_size,
                distance=rest.Distance.COSINE
            )
        )

    def process_text(self, text: str, chunk_size: int = None) -> List[str]:
        """Split text into chunks for embedding."""
        chunk_size = chunk_size or self.config.chunk_size
        overlap = self.config.chunk_overlap
        
        # Simple sentence-based chunking
        sentences = text.split(". ")
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence.split())
            if current_size + sentence_size > chunk_size:
                if current_chunk:
                    chunks.append(". ".join(current_chunk) + ".")
                current_chunk = [sentence]
                current_size = sentence_size
            else:
                current_chunk.append(sentence)
                current_size += sentence_size
        
        if current_chunk:
            chunks.append(". ".join(current_chunk) + ".")
        
        return chunks

    def read_document(self, file_path: str) -> List[DocumentChunk]:
        """Read and process a document file based on its extension."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        if path.suffix not in self.config.supported_extensions:
            raise ValueError(f"Unsupported file type: {path.suffix}")
        
        chunks = []
        metadata = {"source": str(path), "type": path.suffix[1:]}
        
        match path.suffix:
            case ".txt":
                with open(path, 'r') as f:
                    text = f.read()
                    for i, chunk in enumerate(self.process_text(text)):
                        chunks.append(DocumentChunk(
                            text=chunk,
                            metadata=metadata,
                            source_file=str(path),
                            chunk_index=i
                        ))
                        
            case ".md":
                with open(path, 'r') as f:
                    md_text = f.read()
                    html = markdown.markdown(md_text)
                    # Simple HTML to text conversion
                    text = html.replace('<p>', '\\n\\n').replace('</p>', '')
                    for i, chunk in enumerate(self.process_text(text)):
                        chunks.append(DocumentChunk(
                            text=chunk,
                            metadata=metadata,
                            source_file=str(path),
                            chunk_index=i
                        ))
                        
            case ".pdf":
                with open(path, 'rb') as f:
                    pdf = PyPDF2.PdfReader(f)
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() + "\n\n"
                    for i, chunk in enumerate(self.process_text(text)):
                        chunks.append(DocumentChunk(
                            text=chunk,
                            metadata=metadata,
                            source_file=str(path),
                            chunk_index=i
                        ))
                        
            case ".docx":
                doc = docx.Document(path)
                text = "\n\n".join([paragraph.text for paragraph in doc.paragraphs])
                for i, chunk in enumerate(self.process_text(text)):
                    chunks.append(DocumentChunk(
                        text=chunk,
                        metadata=metadata,
                        source_file=str(path),
                        chunk_index=i
                    ))
                    
            case ".json":
                with open(path) as f:
                    data = json.load(f)
                    text = data.get("content", data.get("text", ""))
                    metadata.update(data.get("metadata", {}))
                    for i, chunk in enumerate(self.process_text(text)):
                        chunks.append(DocumentChunk(
                            text=chunk,
                            metadata=metadata,
                            source_file=str(path),
                            chunk_index=i
                        ))
        
        return chunks

    def process_directory(self, directory: str) -> List[DocumentChunk]:
        """Process all supported documents in a directory."""
        path = Path(directory)
        if not path.exists() or not path.is_dir():
            raise NotADirectoryError(f"Directory not found: {directory}")
        
        chunks = []
        for file_path in path.rglob("*"):
            if file_path.suffix in self.config.supported_extensions:
                try:
                    file_chunks = self.read_document(str(file_path))
                    chunks.extend(file_chunks)
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
        
        return chunks

    def create_embeddings(self, chunks: List[DocumentChunk]) -> List[Document]:
        """Create embeddings for document chunks."""
        documents = []
        
        for i in range(0, len(chunks), self.config.batch_size):
            batch = chunks[i:i + self.config.batch_size]
            texts = [chunk.text for chunk in batch]
            
            try:
                response = self.client.embeddings.create(
                    model=self.config.model,
                    input=texts
                )
                
                for chunk, embedding_data in zip(batch, response.data):
                    doc = Document(
                        id=str(uuid.uuid4()),
                        text=chunk.text,
                        metadata={
                            **chunk.metadata,
                            "chunk_index": chunk.chunk_index,
                            "source_file": chunk.source_file
                        },
                        embedding=embedding_data.embedding
                    )
                    documents.append(doc)
                    
            except Exception as e:
                print(f"Error creating embeddings for batch: {e}")
                
        return documents

    def store_documents(self, documents: List[Document]):
        """Store documents in the configured vector database."""
        match self.config.vector_db:
            case QdrantConfig():
                points = [
                    rest.PointStruct(
                        id=doc.id,
                        vector=doc.embedding,
                        payload={
                            "text": doc.text,
                            "metadata": doc.metadata
                        }
                    )
                    for doc in documents
                ]
                self.vector_db.upsert(
                    collection_name=self.config.vector_db.collection_name,
                    points=points
                )
                
            case PineconeConfig():
                vectors = [
                    (
                        doc.id,
                        doc.embedding,
                        {"text": doc.text, "metadata": doc.metadata}
                    )
                    for doc in documents
                ]
                self.vector_db.upsert(vectors=vectors)
                
            case CustomVectorDBConfig():
                if hasattr(self.vector_db, "store_documents"):
                    self.vector_db.store_documents(documents)
                else:
                    raise NotImplementedError(
                        "Custom vector database must implement store_documents method"
                    )

    def search(self, query: str, top_k: int = 5) -> List[Document]:
        """Search for similar documents."""
        # Create query embedding
        query_response = self.client.embeddings.create(
            model=self.config.model,
            input=query
        )
        query_embedding = query_response.data[0].embedding
        
        match self.config.vector_db:
            case QdrantConfig():
                results = self.vector_db.search(
                    collection_name=self.config.vector_db.collection_name,
                    query_vector=query_embedding,
                    limit=top_k
                )
                return [
                    Document(
                        id=str(result.id),
                        text=result.payload["text"],
                        metadata=result.payload["metadata"],
                        embedding=result.vector
                    )
                    for result in results
                ]
                
            case PineconeConfig():
                results = self.vector_db.query(
                    vector=query_embedding,
                    top_k=top_k,
                    include_metadata=True
                )
                return [
                    Document(
                        id=result.id,
                        text=result.metadata["text"],
                        metadata=result.metadata["metadata"],
                        embedding=result.values
                    )
                    for result in results.matches
                ]
                
            case CustomVectorDBConfig():
                if hasattr(self.vector_db, "search"):
                    return self.vector_db.search(query_embedding, top_k)
                else:
                    raise NotImplementedError(
                        "Custom vector database must implement search method"
                    ) 

class DocumentStore:
    """
    Simple interface for document management and search.
    
    Basic usage:
        # Automatically processes all documents in the directory
        docs = DocumentStore("my_documents")
        results = docs.search("my query")
    
    Advanced usage:
        # Manual control over processing
        docs = DocumentStore(
            "my_documents",
            auto_process=False,
            config=EmbeddingConfig(chunk_size=1000)
        )
        docs.add_document("path/to/doc")
        docs.process_batch(batch_size=100)
    """
    
    def __init__(
        self,
        documents_path: Optional[str] = None,
        auto_process: bool = True,
        config: Optional[EmbeddingConfig] = None,
        vector_db_config: Optional[Union[QdrantConfig, PineconeConfig, CustomVectorDBConfig]] = None
    ):
        """
        Initialize document store with optional automatic processing.
        
        Args:
            documents_path: Path to documents directory or file
            auto_process: Whether to process documents automatically
            config: Custom embedding configuration
            vector_db_config: Custom vector database configuration
        """
        self.config = config or EmbeddingConfig()
        if vector_db_config:
            self.config.vector_db = vector_db_config
            
        self.manager = EmbeddingManager(self.config)
        self.processed_docs: List[Document] = []
        
        if documents_path and auto_process:
            self.process_directory(documents_path)
    
    def add_document(self, path: str) -> Optional[Document]:
        """Add and process a single document."""
        try:
            chunks = self.manager.read_document(path)
            docs = self.manager.create_embeddings(chunks)
            self.manager.store_documents(docs)
            self.processed_docs.extend(docs)
            return docs[0] if docs else None
        except Exception as e:
            print(f"Error processing document {path}: {e}")
            return None
    
    def process_directory(self, directory: str, batch_size: Optional[int] = None) -> List[Document]:
        """Process all documents in a directory."""
        try:
            chunks = self.manager.process_directory(directory)
            if batch_size:
                documents = []
                for i in range(0, len(chunks), batch_size):
                    batch = chunks[i:i + batch_size]
                    docs = self.manager.create_embeddings(batch)
                    self.manager.store_documents(docs)
                    documents.extend(docs)
            else:
                documents = self.manager.create_embeddings(chunks)
                self.manager.store_documents(documents)
            
            self.processed_docs.extend(documents)
            return documents
        except Exception as e:
            print(f"Error processing directory {directory}: {e}")
            return []
    
    def search(self, query: str, top_k: int = 3) -> List[Document]:
        """Search for documents similar to the query."""
        return self.manager.search(query, top_k)
    
    def get_document_count(self) -> int:
        """Get the number of processed documents."""
        return len(self.processed_docs)
    
    def clear(self):
        """Clear all processed documents."""
        self.processed_docs = []
        # Recreate collection in vector database
        self.config.recreate_collection = True
        self.manager = EmbeddingManager(self.config) 